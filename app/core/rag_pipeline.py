"""
RAG Pipeline using LangGraph, LangChain, OpenAI Embeddings, and Qdrant.
This file is designed for clarity, modularity, and extensibility.
Brainstorm/TODO comments are included for key decision points.
"""

import os
from dotenv import load_dotenv
import requests
from typing import List, Dict, Any, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI  # Use chat model interface
from langgraph.graph import StateGraph, END, START
import pdfplumber
from docx import Document
from email import policy
from email.parser import BytesParser
from pydantic import BaseModel
from collections import defaultdict
import operator
import threading
import hashlib
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import json
from langchain_qdrant import QdrantVectorStore  # Updated import
from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams
import uuid

# Load environment variables from .env file
load_dotenv()

DOCUMENTS_DIR = "documents"
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

# Get API keys and Qdrant config from environment
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
QDRANT_HOST = os.getenv("QDRANT_HOST", "http://localhost:6333")
QDRANT_PORT = os.getenv("QDRANT_PORT", "6333")
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")

# Performance/env knobs
CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", "800"))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", "100"))
N_TRANSLATIONS_DEFAULT = int(os.getenv("N_TRANSLATIONS", "3"))
TOP_K_DEFAULT = int(os.getenv("TOP_K", "4"))
EMBED_BATCH_SIZE = int(os.getenv("EMBED_BATCH_SIZE", "64"))

# Ensure we have the full URL when a hostname without scheme is provided
if not QDRANT_HOST.startswith("http://") and not QDRANT_HOST.startswith("https://"):
    QDRANT_HOST = f"http://{QDRANT_HOST}"

# Configure logging for timing
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# --- State Schema for LangGraph ---
class RagIngestState(BaseModel):
    load: Optional[List[str]] = None  # document URLs
    chunk: Optional[List[str]] = None  # list of document texts
    embed: Optional[List[Dict[str, Any]]] = None  # list of chunk dicts

class RagQueryState(BaseModel):
    # Remove 'question' from the state to avoid concurrency issues
    translated_queries: Optional[List[str]] = None  # 5 diverse queries
    translation_contexts: Optional[List[Dict[str, Any]]] = None  # Chunks from translation queries
    hyde_context: Optional[List[Dict[str, Any]]] = None  # Chunks from HYDE
    merged_context: Optional[List[Dict[str, Any]]] = None  # Final context for generation
    answer: Optional[Any] = None  # answer from LLM

CACHE_FILE = "embedding_cache.json"

def load_embedding_cache():
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r") as f:
            return json.load(f)
    return {}

def save_embedding_cache(cache):
    with open(CACHE_FILE, "w") as f:
        json.dump(cache, f)

def hash_url(url: str) -> str:
    return hashlib.md5(url.encode()).hexdigest()

def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode('utf-8')).hexdigest()


def _create_qdrant_client_with_fallback() -> QdrantClient:
    """Create a Qdrant client for the configured URL.

    Since we run Qdrant locally (dockerized), prefer the URL from QDRANT_HOST
    and avoid remote/cloud probing or embedded fallbacks.
    """
    timeout_seconds_env = os.getenv("QDRANT_TIMEOUT", "30")
    try:
        timeout_seconds = float(timeout_seconds_env)
    except Exception:
        timeout_seconds = 30.0

    client = QdrantClient(
        url=QDRANT_HOST,
        api_key=QDRANT_API_KEY or None,
        timeout=timeout_seconds,
    )
    return client


def get_qdrant_client() -> QdrantClient:
    """Public helper to obtain a Qdrant client with remote->local fallback."""
    return _create_qdrant_client_with_fallback()

def download_document(url: str, save_dir: str = DOCUMENTS_DIR) -> str:
    # Use a hash of the URL for caching
    filename = hash_url(url) + "_" + url.split("?")[0].split("/")[-1]
    local_filename = os.path.join(save_dir, filename)
    if os.path.exists(local_filename):
        logger.info(f"Cache hit for {url}")
        return local_filename
    logger.info(f"Downloading {url}")
    response = requests.get(url)
    with open(local_filename, "wb") as f:
        f.write(response.content)
    return local_filename

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_text_from_docx(file_path: str) -> str:
    try:
        document = Document(file_path)
        paragraphs = [para.text for para in document.paragraphs if para.text]
        return "\n".join(paragraphs)
    except Exception:
        return ""

def _strip_html_tags(html: str) -> str:
    # Lightweight HTML stripper to avoid extra dependencies
    import re
    # Remove scripts/styles
    html = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
    html = re.sub(r"<style[\s\S]*?</style>", " ", html, flags=re.IGNORECASE)
    # Replace <br> and <p> with newlines
    html = re.sub(r"<(br|/p|/div)>", "\n", html, flags=re.IGNORECASE)
    # Strip all tags
    text = re.sub(r"<[^>]+>", " ", html)
    # Collapse whitespace
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def extract_text_from_eml(file_path: str) -> str:
    try:
        with open(file_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)
        parts_text: list[str] = []
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                try:
                    payload = part.get_content()
                except Exception:
                    payload = None
                if not payload:
                    continue
                if content_type == 'text/plain':
                    parts_text.append(str(payload))
                elif content_type == 'text/html':
                    parts_text.append(_strip_html_tags(str(payload)))
        else:
            content_type = msg.get_content_type()
            payload = msg.get_content()
            if content_type == 'text/plain':
                parts_text.append(str(payload))
            elif content_type == 'text/html':
                parts_text.append(_strip_html_tags(str(payload)))
        return "\n".join(t for t in parts_text if t).strip()
    except Exception:
        return ""

def get_max_threads(n_tasks: int) -> int:
    env_threads = os.getenv("MAX_THREADS")
    if env_threads is not None:
        try:
            env_threads = int(env_threads)
            logger.info(f"Using MAX_THREADS from environment: {env_threads}")
            return min(env_threads, n_tasks)
        except Exception:
            logger.warning(f"Invalid MAX_THREADS value: {env_threads}, falling back to auto-detect.")
    cpu_threads = os.cpu_count() or 8
    max_threads = min(cpu_threads, n_tasks, 16)  # Cap at 16 for safety
    logger.info(f"Using {max_threads} threads for parallel processing (cpu_count={os.cpu_count()})")
    return max_threads

def load_and_preprocess_documents(
    urls: List[str],
    chunk_size: Optional[int] = None,
    chunk_overlap: Optional[int] = None,
    max_threads_override: Optional[int] = None,
) -> List[Dict[str, Any]]:
    start_time = time.time()
    texts = [None] * len(urls)
    file_paths = [None] * len(urls)
    max_threads = max_threads_override if isinstance(max_threads_override, int) and max_threads_override > 0 else get_max_threads(len(urls))
    # Parallel download
    with ThreadPoolExecutor(max_workers=max_threads) as executor:
        future_to_idx = {executor.submit(download_document, url): idx for idx, url in enumerate(urls)}
        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            file_paths[idx] = future.result()
    logger.info(f"Downloaded {len(urls)} documents in {time.time() - start_time:.2f}s")
    # Parallel parse and chunk
    eff_chunk_size = chunk_size if isinstance(chunk_size, int) and chunk_size > 0 else CHUNK_SIZE
    eff_chunk_overlap = chunk_overlap if isinstance(chunk_overlap, int) and chunk_overlap >= 0 else CHUNK_OVERLAP
    def parse_and_chunk(idx, file_path):
        doc_url = urls[idx]
        doc_id = hash_url(doc_url)
        lower_path = file_path.lower()
        if lower_path.endswith(".pdf"):
            # page-aware splitting with metadata
            chunks = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""
                    if not page_text.strip():
                        continue
                    splitter = RecursiveCharacterTextSplitter(chunk_size=eff_chunk_size, chunk_overlap=eff_chunk_overlap)
                    for piece in splitter.split_text(page_text):
                        chunks.append({
                            "id": hash_text(piece),
                            "text": piece,
                            "doc_index": idx,
                            "doc_id": doc_id,
                            "page": page_num,
                        })
            return chunks
        elif lower_path.endswith(".docx"):
            try:
                document = Document(file_path)
                chunks = []
                for p_idx, para in enumerate(document.paragraphs):
                    if not para.text.strip():
                        continue
                    splitter = RecursiveCharacterTextSplitter(chunk_size=eff_chunk_size, chunk_overlap=eff_chunk_overlap)
                    for piece in splitter.split_text(para.text):
                        chunks.append({
                            "id": hash_text(piece),
                            "text": piece,
                            "doc_index": idx,
                            "doc_id": doc_id,
                            "section": f"paragraph_{p_idx+1}",
                        })
                return chunks
            except Exception:
                text = extract_text_from_docx(file_path)
        elif lower_path.endswith(".eml"):
            try:
                with open(file_path, 'rb') as f:
                    msg = BytesParser(policy=policy.default).parse(f)
                subject = msg.get('subject', '')
                sender = msg.get('from', '')
                bodies = []
                if msg.is_multipart():
                    for part in msg.walk():
                        ctype = part.get_content_type()
                        try:
                            payload = part.get_content()
                        except Exception:
                            payload = None
                        if not payload:
                            continue
                        if ctype == 'text/plain':
                            bodies.append(("plain", str(payload)))
                        elif ctype == 'text/html':
                            bodies.append(("html", _strip_html_tags(str(payload))))
                else:
                    ctype = msg.get_content_type()
                    payload = msg.get_content()
                    if ctype == 'text/plain':
                        bodies.append(("plain", str(payload)))
                    elif ctype == 'text/html':
                        bodies.append(("html", _strip_html_tags(str(payload))))
                chunks = []
                splitter = RecursiveCharacterTextSplitter(chunk_size=eff_chunk_size, chunk_overlap=eff_chunk_overlap)
                for part_type, body in bodies:
                    for piece in splitter.split_text(body):
                        chunks.append({
                            "id": hash_text(piece),
                            "text": piece,
                            "doc_index": idx,
                            "doc_id": doc_id,
                            "part": part_type,
                            "subject": subject,
                            "from": sender,
                        })
                return chunks
            except Exception:
                text = extract_text_from_eml(file_path)
        elif lower_path.endswith(".txt"):
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception:
                text = ""
        else:
            text = ""
        splitter = RecursiveCharacterTextSplitter(chunk_size=eff_chunk_size, chunk_overlap=eff_chunk_overlap)
        chunks = []
        for piece in splitter.split_text(text):
            chunks.append({
                "id": hash_text(piece),
                "text": piece,
                "doc_index": idx,
                "doc_id": doc_id,
            })
        return chunks
    all_chunks = []
    with ThreadPoolExecutor(max_workers=max_threads) as executor:
        future_to_idx = {executor.submit(parse_and_chunk, idx, file_path): idx for idx, file_path in enumerate(file_paths)}
        for future in as_completed(future_to_idx):
            chunks = future.result()
            all_chunks.extend(chunks)
    logger.info(f"Parsed and chunked {len(file_paths)} documents into {len(all_chunks)} chunks in {time.time() - start_time:.2f}s (total)")
    return all_chunks

def chunk_documents(texts: List[str]) -> List[Dict[str, Any]]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = []
    for idx, text in enumerate(texts):
        for chunk in splitter.split_text(text):
            chunks.append({"text": chunk, "doc_index": idx})
    return chunks

def embed_and_store(chunks: List[Dict[str, Any]], collection_name: str = "docs") -> QdrantVectorStore:
    import time
    embeddings_model = OpenAIEmbeddings(
        openai_api_key=OPENAI_API_KEY,
        model="text-embedding-3-large"
    )
    
    # Check if caching is enabled
    enable_qdrant_cache = os.getenv("ENABLE_QDRANT_CACHE", "true").lower() == "true"
    
    # Prepare chunk data
    texts = [c["text"] for c in chunks]
    metadatas = []
    ids = []
    for c in chunks:
        meta = {"doc_index": c.get("doc_index"), "doc_id": c.get("doc_id")}
        if "page" in c:
            meta["page"] = c["page"]
        if "section" in c:
            meta["section"] = c["section"]
        if "part" in c:
            meta["part"] = c["part"]
        if "subject" in c:
            meta["subject"] = c["subject"]
        if "from" in c:
            meta["from"] = c["from"]
        metadatas.append(meta)
        # Generate a stable UUIDv5 for Qdrant point ID
        salt = f"{c.get('doc_id','')}|{meta.get('page') or meta.get('section') or meta.get('part') or ''}|{hash_text(c['text'])}"
        point_uuid = str(uuid.uuid5(uuid.NAMESPACE_URL, salt))
        ids.append(point_uuid)
    
    start_time = time.time()
    logger.info(f"Processing {len(texts)} chunks for Qdrant...")
    
    # QdrantClient approach (robust for all versions)
    try:
        client = _create_qdrant_client_with_fallback()

        # Ensure collection exists (retry a few times on transient failures)
        max_retries = int(os.getenv("QDRANT_RETRIES", "3"))
        collection_exists = False
        for attempt in range(1, max_retries + 1):
            try:
                try:
                    collection_info = client.get_collection(collection_name=collection_name)
                    logger.info(f"Collection '{collection_name}' already exists with {collection_info.points_count} points")
                    collection_exists = True
                except Exception:
                    logger.info(f"Creating collection '{collection_name}' with vector size 3072")
                    client.create_collection(
                        collection_name=collection_name,
                        vectors_config=VectorParams(size=3072, distance=Distance.COSINE),
                    )
                    collection_exists = True
                break
            except Exception as e:
                if attempt == max_retries:
                    raise
                wait_s = min(2 ** attempt, 10)
                logger.warning(f"Collection ensure failed (attempt {attempt}/{max_retries}): {e}. Retrying in {wait_s}s...")
                time.sleep(wait_s)

        vectordb = QdrantVectorStore(
            client=client,
            collection_name=collection_name,
            embedding=embeddings_model,
        )

        # Check for existing chunks if caching is enabled and collection exists
        new_chunks_indices = list(range(len(chunks)))
        if enable_qdrant_cache and collection_exists:
            try:
                # Get existing point IDs in batches to avoid memory issues
                existing_ids = set()
                offset = None
                batch_size = 1000
                
                while True:
                    scroll_result = client.scroll(
                        collection_name=collection_name,
                        limit=batch_size,
                        offset=offset,
                        with_payload=False,
                        with_vectors=False
                    )
                    points, next_offset = scroll_result
                    
                    if not points:
                        break
                        
                    existing_ids.update(point.id for point in points)
                    offset = next_offset
                    
                    if next_offset is None:
                        break
                
                logger.info(f"Found {len(existing_ids)} existing chunks in Qdrant")
                
                # Filter out chunks that already exist
                new_chunks_indices = []
                for i, chunk_id in enumerate(ids):
                    if chunk_id not in existing_ids:
                        new_chunks_indices.append(i)
                
                if not new_chunks_indices:
                    logger.info("All chunks already exist in Qdrant, skipping embedding and storage")
                    return vectordb
                
                logger.info(f"Adding {len(new_chunks_indices)} new chunks (skipped {len(chunks) - len(new_chunks_indices)} existing)")
                
            except Exception as e:
                logger.warning(f"Could not check existing chunks, proceeding with full upsert: {e}")
                new_chunks_indices = list(range(len(chunks)))

        # Prepare data for new chunks only
        if new_chunks_indices != list(range(len(chunks))):
            texts = [texts[i] for i in new_chunks_indices]
            metadatas = [metadatas[i] for i in new_chunks_indices]
            ids = [ids[i] for i in new_chunks_indices]

        if texts:  # Only proceed if there are chunks to add
            # Upsert with retries (idempotent via ids)
            for attempt in range(1, max_retries + 1):
                try:
                    vectordb.add_texts(texts=texts, metadatas=metadatas, ids=ids)
                    logger.info(
                        f"Embedded and upserted {len(texts)} chunks to Qdrant in {time.time() - start_time:.2f}s"
                    )
                    break
                except Exception as e:
                    if attempt == max_retries:
                        raise
                    wait_s = min(2 ** attempt, 10)
                    logger.warning(f"Upsert failed (attempt {attempt}/{max_retries}): {e}. Retrying in {wait_s}s...")
                    time.sleep(wait_s)
        else:
            logger.info(f"No new chunks to add, completed in {time.time() - start_time:.2f}s")
            
    except Exception as e:
        logger.error(f"Error during Qdrant upsert: {e}", exc_info=True)
        raise
    return vectordb

def retrieve_context(vectordb: QdrantVectorStore, query: str, k: int = TOP_K_DEFAULT) -> List[Dict[str, Any]]:
    try:
        doc_scores = vectordb.similarity_search_with_score(query, k=k)
        scores = [s for _, s in doc_scores]
        if scores:
            s_min, s_max = min(scores), max(scores)
            denom = (s_max - s_min) or 1.0
        else:
            s_min, denom = 0.0, 1.0
        results = []
        for rank, (doc, score) in enumerate(doc_scores, start=1):
            results.append({
                "text": doc.page_content,
                "metadata": doc.metadata,
                "score": score,
                "score_normalized": (score - s_min) / denom,
                "rank": rank,
            })
        return results
    except Exception:
        docs = vectordb.similarity_search(query, k=k)
        return [{"text": d.page_content, "metadata": d.metadata, "score": getattr(d, 'score', None), "rank": i+1} for i, d in enumerate(docs)]

def generate_answer(context: List[Dict[str, Any]], query: str) -> str:
    # Use OpenAI's GPT-4o for final grounded answer generation via chat completions
    llm = ChatOpenAI(
        temperature=0,
        openai_api_key=OPENAI_API_KEY,
        model="gpt-4o"
    )
    context_str = "\n\n".join([c["text"] for c in context])
    prompt = (
        "You are an expert insurance policy analyst. Extract the exact answer from the context below.\n"
        "Return a JSON object with keys: decision (Yes/No/Unclear), answer, rationale, quotes.\n"
        "- decision: Yes if context clearly supports; No if it clearly denies; Unclear otherwise.\n"
        "- answer: Write a professional, complete response that includes ALL specific numbers, percentages, time periods, conditions, and requirements from the context. Use exact figures (e.g., 'thirty (30) days', 'thirty-six (36) months', '1% of Sum Insured', 'two (2) years'). Start with 'Yes,' or 'No,' when applicable. Be comprehensive but concise - include eligibility criteria, limits, exclusions, and procedural requirements as stated in the policy.\n"
        "- rationale: Brief explanation of the evidence.\n"
        "- quotes: 1-2 key verbatim excerpts that support your answer.\n"
        "If insufficient evidence exists, mark decision=Unclear and explain what information is missing.\n\n"
        f"Context:\n{context_str}\n\n"
        f"Question: {query}\n"
        "JSON:"
    )
    msg = llm.invoke(prompt)
    try:
        return msg.content  # ChatOpenAI returns an AIMessage
    except AttributeError:
        return str(msg)

def format_output(answer: str, context: List[Dict[str, Any]]) -> Dict[str, Any]:
    return {
        "answer": answer,
        "supporting_clauses": [
            {
                "text": c["text"],
                "doc_index": c["metadata"].get("doc_index"),
                "score": c.get("score"),
            }
            for c in context
        ],
        "rationale": "See supporting clauses and cited text for reasoning."
    }

# --- Reciprocal Rank Fusion (RRF) Algorithm ---
def reciprocal_rank_fusion(*list_of_list_ranks_system, K=60):
    """
    Fuse rank from multiple IR systems using Reciprocal Rank Fusion.
    Args:
    * list_of_list_ranks_system: Ranked results from different IR system.
    K (int): A constant used in the RRF formula (default is 60).
    Returns:
    Tuple of list of sorted documents by score and sorted documents
    """
    rrf_map = defaultdict(float)
    for rank_list in list_of_list_ranks_system:
        for rank, item in enumerate(rank_list, 1):
            # Use a tuple of (text, doc_index) as a unique key for deduplication
            key = (item['text'], str(item['metadata'].get('doc_index', '')))
            rrf_map[key] += 1 / (rank + K)
    # Map back to original chunk objects
    chunk_lookup = { (item['text'], str(item['metadata'].get('doc_index', ''))): item for rank_list in list_of_list_ranks_system for item in rank_list }
    sorted_items = sorted(rrf_map.items(), key=lambda x: x[1], reverse=True)
    sorted_chunks = [chunk_lookup[key] for key, score in sorted_items]
    return sorted_items, sorted_chunks

# --- Production-Grade Prompts ---
QUERY_TRANSLATION_PROMPT = (
    """
    You are an expert legal/insurance/HR assistant. Given the following user question, generate 5 diverse queries that:
    - Rephrase the original question in different ways
    - Widen the scope to cover related aspects
    - Break down compound or multi-part questions into focused sub-questions
    - Use different perspectives or approaches to maximize information retrieval
    Each query should be clear, information-seeking, and non-redundant. Return the 5 queries as a numbered list.
    
    User question: {question}
    """
)

HYDE_PROMPT = (
    """
    You are a world-class legal/insurance/HR expert. Given the following user question, generate a concise, fact-rich, and highly relevant hypothetical answer as if you had access to all authoritative documents. Your answer should be clear, actionable, and maximize the value for a professional user. Avoid speculation and focus on likely, useful information.
    
    User question: {question}
    """
)

# --- Ingestion Pipeline Nodes ---
def ingest_node_load(state: RagIngestState) -> RagIngestState:
    # Now returns all chunks directly
    chunks = load_and_preprocess_documents(state.load)
    return state.copy(update={"embed": chunks})

def ingest_node_chunk(state: RagIngestState) -> RagIngestState:
    # No-op, as chunking is now done in load step
    return state

def ingest_node_embed(state: RagIngestState, collection_name: str = "docs") -> RagIngestState:
    embed_and_store(state.embed, collection_name=collection_name)
    return state  # No update needed, as storage is side-effect

# --- Ingestion Pipeline Graph ---
def build_ingest_graph(collection_name: str = "docs", *, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None, max_threads: Optional[int] = None):
    graph = StateGraph(RagIngestState)
    def load_node(state: RagIngestState) -> RagIngestState:
        chunks = load_and_preprocess_documents(state.load, chunk_size=chunk_size, chunk_overlap=chunk_overlap, max_threads_override=max_threads)
        return state.copy(update={"embed": chunks})
    graph.add_node("load", load_node)
    graph.add_node("chunk", ingest_node_chunk)
    # Use a lambda to pass collection_name
    graph.add_node("embed", lambda state: ingest_node_embed(state, collection_name=collection_name))
    graph.add_edge(START, "load")
    graph.add_edge("load", "chunk")
    graph.add_edge("chunk", "embed")
    graph.add_edge("embed", END)
    return graph.compile()

# --- Query Pipeline Nodes ---
def query_node_translate(state: RagQueryState, question, small_llm, vectordb, n: int = 5, k: int = 5) -> RagQueryState:
    prompt = QUERY_TRANSLATION_PROMPT.format(question=question)
    llm_output = small_llm.invoke(prompt)
    output_text = getattr(llm_output, "content", str(llm_output))
    translated_queries = [q.strip(" .-") for q in output_text.split("\n") if q.strip()][:n]
    all_chunks_per_query = []
    for tq in translated_queries:
        chunks = retrieve_context(vectordb, tq, k=k)
        for chunk in chunks:
            chunk['source_query'] = tq
        all_chunks_per_query.append(chunks)
    _, ranked_chunks = reciprocal_rank_fusion(*all_chunks_per_query, K=60)
    ranked_chunks = ranked_chunks[:k*n]
    return state.copy(update={
        "translated_queries": translated_queries,
        "translation_contexts": ranked_chunks
    })

def query_node_hyde(state: RagQueryState, question, gpt4_llm) -> RagQueryState:
    prompt = HYDE_PROMPT.format(question=question)
    hyde_msg = gpt4_llm.invoke(prompt)
    hyde_answer = getattr(hyde_msg, "content", str(hyde_msg))
    hyde_context = [{"text": hyde_answer, "metadata": {"source": "HYDE"}}]
    return state.copy(update={"hyde_context": hyde_context})

def query_node_merge(state: RagQueryState) -> RagQueryState:
    merged = (state.translation_contexts or []) + (state.hyde_context or [])
    return state.copy(update={"merged_context": merged})

def query_node_generate(state: RagQueryState, question) -> RagQueryState:
    answer = generate_answer(state.merged_context, question)
    return state.copy(update={"answer": answer})

# --- Query Pipeline Graph ---
def build_query_graph(vectordb, small_llm, gpt4_llm, n: int = N_TRANSLATIONS_DEFAULT, k: int = TOP_K_DEFAULT, question=None):
    def translate_node(state):
        return query_node_translate(state, question, small_llm, vectordb, n, k)
    def hyde_node(state):
        return query_node_hyde(state, question, gpt4_llm)
    def generate_node(state):
        return query_node_generate(state, question)
    graph = StateGraph(RagQueryState)
    graph.add_node("translate", translate_node)
    graph.add_node("hyde", hyde_node)
    graph.add_node("merge", query_node_merge)
    graph.add_node("generate", generate_node)
    graph.add_edge(START, "translate")
    graph.add_edge(START, "hyde")
    graph.add_edge("translate", "merge")
    graph.add_edge("hyde", "merge")
    graph.add_edge("merge", "generate")
    graph.add_edge("generate", END)
    return graph.compile()

# --- Entry Points ---
def ingest_documents(document_urls: List[str], collection_name: str = "docs", *, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None, max_threads: Optional[int] = None):
    graph = build_ingest_graph(collection_name, chunk_size=chunk_size, chunk_overlap=chunk_overlap, max_threads=max_threads)
    state = RagIngestState(load=document_urls)
    graph.invoke(state)
    # No return needed; data is stored in Qdrant

def answer_questions(questions: List[str], vectordb, k: int = 5) -> Dict[str, Any]:
    graph = build_query_graph(vectordb, k)
    answers = []
    for question in questions:
        state = RagQueryState()
        result = graph.invoke(state)
        answers.append(result["answer"])
    return {"answers": answers}

def answer_questions_advanced(questions: List[str], vectordb, small_llm, gpt4_llm, n: int = 5, k: int = 5) -> Dict[str, Any]:
    answers = []
    detailed = []
    for question in questions:
        graph = build_query_graph(vectordb, small_llm, gpt4_llm, n, k, question=question)
        state = RagQueryState()
        result = graph.invoke(state)
        # result["answer"] contains the JSON string from generate_answer
        answers.append(result["answer"])
        # construct a detailed object with citations if available
        citations = []
        for ctx in (state.translation_contexts or [])[:k]:
            meta = ctx.get("metadata", {})
            citations.append({
                "doc_index": meta.get("doc_index"),
                "page": meta.get("page"),
                "section": meta.get("section"),
                "part": meta.get("part"),
                "score": ctx.get("score_normalized", ctx.get("score")),
                "snippet": ctx.get("text"),
                "source_query": ctx.get("source_query"),
            })
        detailed.append({
            "question": question,
            "model_json": result["answer"],
            "citations": citations,
            "used_queries": state.translated_queries or [],
        })
    return {"answers": answers, "answers_detailed": detailed}

# --- Example Usage ---
if __name__ == "__main__":
    document_urls = [
        "https://hackrx.blob.core.windows.net/assets/policy.pdf?..."
    ]
    questions = [
        "What is the grace period for premium payment under the National Parivar Mediclaim Plus Policy?",
        "What is the waiting period for pre-existing diseases (PED) to be covered?"
    ]
    print("Starting document ingestion...")
    ingest_documents(document_urls, collection_name="docs")
    print("Ingestion complete.")
    print("Initializing Qdrant client and vector store for querying...")
    embeddings_model = OpenAIEmbeddings(
        openai_api_key=OPENAI_API_KEY,
        model="text-embedding-3-large"
    )
    client = QdrantClient(
        url=QDRANT_HOST,
        api_key=QDRANT_API_KEY or None,
    )
    vectordb = QdrantVectorStore(
        client=client,
        collection_name="docs",
        embedding=embeddings_model,
    )
    print("Answering questions...")
    from langchain_openai import ChatOpenAI
    small_llm = ChatOpenAI(temperature=0, openai_api_key=OPENAI_API_KEY, model="gpt-4o-mini")
    gpt4_llm = ChatOpenAI(temperature=0, openai_api_key=OPENAI_API_KEY, model="gpt-4o-mini")
    output = answer_questions_advanced(questions, vectordb, small_llm, gpt4_llm)
    print("\n--- Final Output ---")
    import json
    print(json.dumps(output, indent=2))